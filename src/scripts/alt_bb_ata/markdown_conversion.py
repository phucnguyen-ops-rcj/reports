from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont, ImageOps
from PIL import JpegImagePlugin  # noqa: F401

IMAGE_PATTERN = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)$")

PAGE_WIDTH = 1240
PAGE_HEIGHT = 1754
PAGE_MARGIN_X = 72
PAGE_MARGIN_Y = 72
LINE_SPACING = 12
PARAGRAPH_SPACING = 18
SECTION_SPACING = 24
IMAGE_SPACING = 20


@dataclass(slots=True)
class RenderState:
    pages: list[Image.Image]
    page: Image.Image
    draw: ImageDraw.ImageDraw
    y: int


def markdown_to_pdf(
    markdown_path: str | Path,
    output_path: str | Path | None = None,
) -> Path:
    markdown_file = Path(markdown_path)
    if not markdown_file.exists():
        raise FileNotFoundError(f"Markdown file not found: {markdown_file}")

    pdf_path = _resolve_output_path(markdown_file, output_path, ".pdf")
    lines = _read_markdown_lines(markdown_file)

    fonts = _load_fonts()
    state = _new_state()

    for raw_line in lines:
        line = raw_line.rstrip()
        image_match = IMAGE_PATTERN.match(line)
        if image_match:
            _draw_image_block(
                state,
                markdown_file.parent / image_match.group("path"),
                alt_text=image_match.group("alt") or "image",
                fonts=fonts,
            )
            continue

        if not line.strip():
            state.y += PARAGRAPH_SPACING
            continue

        if line.startswith("# "):
            _draw_text_block(
                state,
                line[2:],
                fonts["h1"],
                fill="black",
                extra_spacing=SECTION_SPACING,
            )
            continue
        if line.startswith("## "):
            _draw_text_block(
                state,
                _strip_html_underline(line[3:]),
                fonts["h2"],
                fill="black",
                extra_spacing=SECTION_SPACING,
            )
            continue
        if line.startswith("### "):
            _draw_text_block(
                state,
                _strip_html_underline(line[4:]),
                fonts["h3"],
                fill="black",
                extra_spacing=PARAGRAPH_SPACING,
            )
            continue
        if line.startswith("- "):
            _draw_bullet_block(state, line[2:], fonts)
            continue
        if line.startswith("http://") or line.startswith("https://"):
            _draw_text_block(
                state,
                line,
                fonts["link"],
                fill="#0b57d0",
                extra_spacing=PARAGRAPH_SPACING,
            )
            continue

        _draw_text_block(
            state,
            _strip_markdown_emphasis(line),
            fonts["body"],
            fill="black",
            extra_spacing=PARAGRAPH_SPACING,
        )

    state.pages.append(state.page)
    _save_pdf(state.pages, pdf_path)
    return pdf_path


def markdown_to_word(
    markdown_path: str | Path,
    output_path: str | Path | None = None,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for Word export. Install project dependencies first."
        ) from exc

    markdown_file = Path(markdown_path)
    if not markdown_file.exists():
        raise FileNotFoundError(f"Markdown file not found: {markdown_file}")

    word_path = _resolve_output_path(markdown_file, output_path, ".docx")
    lines = _read_markdown_lines(markdown_file)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    body_style = document.styles["Normal"]
    body_style.font.name = "Arial"
    body_style.font.size = Pt(11)

    max_image_width = section.page_width - section.left_margin - section.right_margin

    for raw_line in lines:
        line = raw_line.rstrip()
        image_match = IMAGE_PATTERN.match(line)
        if image_match:
            image_path = markdown_file.parent / image_match.group("path")
            if image_path.exists():
                document.add_picture(str(image_path), width=max_image_width)
            else:
                document.add_paragraph(
                    f"[Missing image: {image_path.name} ({image_match.group('alt') or 'image'})]"
                )
            continue

        if not line.strip():
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
            continue
        if line.startswith("## "):
            document.add_heading(_strip_html_underline(line[3:]), level=2)
            continue
        if line.startswith("### "):
            document.add_heading(_strip_html_underline(line[4:]), level=3)
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(_strip_markdown_emphasis(line[2:]))
            continue

        paragraph = document.add_paragraph()
        run = paragraph.add_run(_strip_markdown_emphasis(line))
        if line.startswith("http://") or line.startswith("https://"):
            font = run.font
            font.color.rgb = _rgb_color("0B57D0")
            font.underline = True

    word_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(word_path)
    return word_path


def _new_state() -> RenderState:
    page = Image.new("RGB", (PAGE_WIDTH, PAGE_HEIGHT), "white")
    return RenderState(
        pages=[],
        page=page,
        draw=ImageDraw.Draw(page),
        y=PAGE_MARGIN_Y,
    )


def _load_fonts() -> dict[str, ImageFont.ImageFont]:
    font_candidates = [
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    mono_candidates = [
        "/System/Library/Fonts/SFNSMono.ttf",
        "/System/Library/Fonts/Supplemental/Courier New.ttf",
    ]

    regular = _load_font(font_candidates[0], 28)
    bold = _load_font(font_candidates[1], 28)
    return {
        "h1": _load_font(font_candidates[1], 34),
        "h2": _load_font(font_candidates[1], 28),
        "h3": _load_font(font_candidates[1], 24),
        "body": regular,
        "bullet": regular,
        "link": regular,
        "caption": _load_font(font_candidates[0], 22),
        "mono": _load_font(mono_candidates[0], 24),
        "bold": bold,
    }


def _load_font(path: str, size: int) -> ImageFont.ImageFont:
    try:
        return ImageFont.truetype(path, size=size)
    except Exception:
        return ImageFont.load_default()


def _draw_text_block(
    state: RenderState,
    text: str,
    font: ImageFont.ImageFont,
    *,
    fill: str,
    extra_spacing: int,
    indent: int = 0,
) -> None:
    wrapped_lines = _wrap_text(
        state.draw, text, font, PAGE_WIDTH - (2 * PAGE_MARGIN_X) - indent
    )
    line_height = _line_height(font)
    required_height = len(wrapped_lines) * (line_height + LINE_SPACING) + extra_spacing
    _ensure_space(state, required_height)

    x = PAGE_MARGIN_X + indent
    for line in wrapped_lines:
        state.draw.text((x, state.y), line, font=font, fill=fill)
        state.y += line_height + LINE_SPACING
    state.y += extra_spacing


def _draw_bullet_block(
    state: RenderState,
    text: str,
    fonts: dict[str, ImageFont.ImageFont],
) -> None:
    bullet_indent = 36
    text_indent = 64
    wrapped_lines = _wrap_text(
        state.draw,
        _strip_markdown_emphasis(text),
        fonts["bullet"],
        PAGE_WIDTH - (2 * PAGE_MARGIN_X) - text_indent,
    )
    line_height = _line_height(fonts["bullet"])
    required_height = (
        len(wrapped_lines) * (line_height + LINE_SPACING) + PARAGRAPH_SPACING
    )
    _ensure_space(state, required_height)

    state.draw.text(
        (PAGE_MARGIN_X + bullet_indent, state.y),
        "\u2022",
        font=fonts["bullet"],
        fill="black",
    )
    for index, line in enumerate(wrapped_lines):
        x = PAGE_MARGIN_X + text_indent
        y = state.y + (index * (line_height + LINE_SPACING))
        state.draw.text((x, y), line, font=fonts["bullet"], fill="black")

    state.y += len(wrapped_lines) * (line_height + LINE_SPACING) + PARAGRAPH_SPACING


def _draw_image_block(
    state: RenderState,
    image_path: Path,
    *,
    alt_text: str,
    fonts: dict[str, ImageFont.ImageFont],
) -> None:
    image = _open_image_or_placeholder(image_path, alt_text, fonts["caption"])
    max_width = PAGE_WIDTH - (2 * PAGE_MARGIN_X)
    max_height = PAGE_HEIGHT - (2 * PAGE_MARGIN_Y) - 160
    image.thumbnail((max_width, max_height))

    caption = image_path.name
    caption_lines = _wrap_text(state.draw, caption, fonts["caption"], max_width)
    caption_height = len(caption_lines) * (_line_height(fonts["caption"]) + 8)
    required_height = image.height + IMAGE_SPACING + caption_height + PARAGRAPH_SPACING
    _ensure_space(state, required_height)

    x = PAGE_MARGIN_X + ((max_width - image.width) // 2)
    state.page.paste(image, (x, state.y))
    state.y += image.height + IMAGE_SPACING

    for line in caption_lines:
        state.draw.text(
            (PAGE_MARGIN_X, state.y), line, font=fonts["caption"], fill="#555555"
        )
        state.y += _line_height(fonts["caption"]) + 8
    state.y += PARAGRAPH_SPACING


def _open_image_or_placeholder(
    image_path: Path,
    alt_text: str,
    caption_font: ImageFont.ImageFont,
) -> Image.Image:
    if image_path.exists():
        with Image.open(image_path) as image:
            return image.convert("RGB")

    placeholder = Image.new("RGB", (900, 520), "#f3f3f3")
    draw = ImageDraw.Draw(placeholder)
    draw.rectangle((0, 0, 899, 519), outline="#cccccc", width=3)
    lines = wrap(f"Missing image: {image_path.name} ({alt_text})", width=40)
    y = 220
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=caption_font)
        x = (placeholder.width - (bbox[2] - bbox[0])) // 2
        draw.text((x, y), line, font=caption_font, fill="#666666")
        y += _line_height(caption_font) + 10
    return placeholder


def _ensure_space(state: RenderState, required_height: int) -> None:
    bottom_limit = PAGE_HEIGHT - PAGE_MARGIN_Y
    if state.y + required_height <= bottom_limit:
        return
    state.pages.append(state.page)
    next_state = _new_state()
    state.page = next_state.page
    state.draw = next_state.draw
    state.y = next_state.y


def _wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.ImageFont,
    width: int,
) -> list[str]:
    words = text.split()
    if not words:
        return [""]

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textlength(candidate, font=font) <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _line_height(font: ImageFont.ImageFont) -> int:
    bbox = font.getbbox("Ag")
    return bbox[3] - bbox[1]


def _read_markdown_lines(markdown_file: Path) -> list[str]:
    return markdown_file.read_text(encoding="utf-8").splitlines()


def _resolve_output_path(
    markdown_file: Path,
    output_path: str | Path | None,
    suffix: str,
) -> Path:
    return (
        Path(output_path)
        if output_path is not None
        else markdown_file.with_suffix(suffix)
    )


def _strip_html_underline(text: str) -> str:
    return text.replace("<u>", "").replace("</u>", "")


def _strip_markdown_emphasis(text: str) -> str:
    return text.replace("**", "")


def _rgb_color(hex_value: str):
    from docx.shared import RGBColor

    return RGBColor.from_string(hex_value)


def _save_pdf(pages: list[Image.Image], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    first, *rest = [ImageOps.exif_transpose(page).convert("RGB") for page in pages]
    first.save(output_path, save_all=True, append_images=rest, resolution=150.0)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert an ALT/BB/ATA markdown report to PDF or Word."
    )
    parser.add_argument("markdown_path", help="Path to the markdown report.")
    parser.add_argument(
        "--format",
        choices=("pdf", "docx"),
        default="pdf",
        help="Output format. Defaults to pdf.",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Output file path. Defaults to the markdown path with the selected suffix.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.format == "docx":
        output_path = markdown_to_word(args.markdown_path, args.output)
    else:
        output_path = markdown_to_pdf(args.markdown_path, args.output)
    print(output_path)


if __name__ == "__main__":
    main()
